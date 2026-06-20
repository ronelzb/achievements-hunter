"""DOCX renderer and user-annotation extractor for steam-platinum."""

from __future__ import annotations

import difflib
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor

from .contracts import (
    LLMStrategyOutput,
    PendingAchievement,
    StrategyResult,
    UserAnnotations,
)

_CATEGORY_COLOR: dict[str, tuple[int, int, int]] = {
    "missable": (0xFF, 0x00, 0x00),  # red
    "story": (0x44, 0x72, 0xC4),  # steel blue
    "grind": (0xE3, 0x6C, 0x09),  # orange
    "collectible": (0xE3, 0x6C, 0x09),  # orange
    "difficulty": (0x70, 0x30, 0xA0),  # purple
    "misc": (0x40, 0x40, 0x40),  # dark grey
}

_HYPERLINK_REL = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"
)


def _para_full_text(para) -> str:
    """Return all text from a paragraph including text inside hyperlink elements."""
    parts: list[str] = []
    for elem in para._p:
        local = elem.tag.split("}", 1)[-1] if "}" in elem.tag else elem.tag
        if local == "r":
            parts.extend(t.text for t in elem.iter(qn("w:t")) if t.text)
        elif local == "hyperlink":
            for child in elem:
                child_local = (
                    child.tag.split("}", 1)[-1] if "}" in child.tag else child.tag
                )
                if child_local == "r":
                    parts.extend(t.text for t in child.iter(qn("w:t")) if t.text)
    return "".join(parts)


def _add_hyperlink(para, text: str, url: str) -> None:
    """Append a Hyperlink-styled run to *para*."""
    r_id = para.part.relate_to(url, _HYPERLINK_REL, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_style = OxmlElement("w:rStyle")
    r_style.set(qn("w:val"), "Hyperlink")
    r_pr.append(r_style)
    run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    para._p.append(hyperlink)


def render_docx(
    result: StrategyResult,
    pending: list[PendingAchievement],
    path: Path,
) -> None:
    """Render *result* to a DOCX file at *path*."""
    doc = Document()
    output = result.output
    date_str = result.created_at.strftime("%Y-%m-%d")

    doc.add_heading(f"{result.game_name} — Platinum Guide", level=1)
    doc.add_paragraph(
        f"Generated: {date_str}  |  {result.model}  |  Pending: {len(pending)}"
    )

    summary_para = doc.add_paragraph(output.summary)
    for run in summary_para.runs:
        run.font.italic = True

    runs_para = doc.add_paragraph(
        f"Minimum runs: {output.total_runs}  •  Est. hours: {output.estimated_hours}"
    )
    for run in runs_para.runs:
        run.font.bold = True

    for section in output.sections:
        r, g, b = _CATEGORY_COLOR.get(section.category, (0x40, 0x40, 0x40))
        heading = doc.add_heading(section.title, level=2)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(r, g, b)

        doc.add_paragraph(section.overview)

        for item in section.items:
            item_para = doc.add_paragraph()
            item_para.add_run("☐  ")
            if item.guide_link:
                _add_hyperlink(item_para, item.display_name, item.guide_link)
            else:
                item_para.add_run(item.display_name)
            doc.add_paragraph(f"     {item.tip}")

    doc.add_heading("Recommended Order", level=2)
    for i, step in enumerate(output.recommended_order, start=1):
        doc.add_paragraph(f"  {i:2}. {step}")

    doc.add_heading("My Notes", level=2)
    doc.add_paragraph(
        "Add notes anywhere in this document — tips that didn't work, "
        "things to try, questions. Run --refine to incorporate them."
    )

    doc.add_paragraph(f"steam-platinum v1  •  App ID {result.app_id}")
    doc.save(str(path))


def render_to_text(output: LLMStrategyOutput) -> str:
    """Render *output* to plain text for use as the baseline in --refine diffs."""
    lines: list[str] = [
        output.summary,
        f"Minimum runs: {output.total_runs}  •  Est. hours: {output.estimated_hours}",
    ]
    for section in output.sections:
        lines.append(section.title)
        lines.append(section.overview)
        for item in section.items:
            lines.append(f"☐  {item.display_name}")
            lines.append(f"     {item.tip}")
    lines.append("Recommended Order")
    for i, step in enumerate(output.recommended_order, start=1):
        lines.append(f"  {i:2}. {step}")
    lines.append("My Notes")
    lines.append(
        "Add notes anywhere in this document — tips that didn't work, "
        "things to try, questions. Run --refine to incorporate them."
    )
    return "\n".join(lines)


def extract_annotations(path: Path, baseline_text: str) -> UserAnnotations:
    """Diff the DOCX at *path* against *baseline_text* and return the delta."""
    doc = Document(str(path))
    docx_lines = [line for p in doc.paragraphs if (line := _para_full_text(p).strip())]
    baseline_lines = [line for line in baseline_text.splitlines() if line.strip()]

    added: list[str] = []
    modified: list[str] = []
    deleted: list[str] = []

    sm = difflib.SequenceMatcher(None, baseline_lines, docx_lines, autojunk=False)
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == "replace":
            modified.extend(docx_lines[j1:j2])
        elif tag == "insert":
            added.extend(docx_lines[j1:j2])
        elif tag == "delete":
            deleted.extend(baseline_lines[i1:i2])

    return UserAnnotations(
        docx_path=str(path),
        added_lines=added,
        modified_lines=modified,
        deleted_lines=deleted,
    )
