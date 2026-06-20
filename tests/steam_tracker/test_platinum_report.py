"""Tests for platinum_report.py.

render_docx is tested by writing a real DOCX to a tmp_path and reading it back.
extract_annotations is tested by constructing a DOCX with known content and diffing.
"""

from __future__ import annotations

from datetime import UTC, datetime
from pathlib import Path

from docx import Document

from steam_tracker.contracts import (
    LLMStrategyOutput,
    PendingAchievement,
    StrategyItem,
    StrategyResult,
    StrategySection,
)
from steam_tracker.platinum_report import (
    _para_full_text,
    extract_annotations,
    render_docx,
    render_to_text,
)

# ── Fixtures ──────────────────────────────────────────────────────────────────


def _make_output() -> LLMStrategyOutput:
    return LLMStrategyOutput(
        total_runs=2,
        estimated_hours="20-30",
        summary="Complete on Survival for all story achievements.",
        sections=[
            StrategySection(
                title="Missable Achievements",
                category="missable",
                overview="These can be missed on a playthrough.",
                items=[
                    StrategyItem(
                        api_name="ACH_1",
                        display_name="Don't Look Back",
                        tip="Avoid looking at enemies in Chapter 3.",
                        guide_link="https://www.truesteamachievements.com/a/dont-look-back",
                    ),
                ],
            ),
            StrategySection(
                title="Story Achievements",
                category="story",
                overview="Earned naturally through the story.",
                items=[
                    StrategyItem(
                        api_name="ACH_2",
                        display_name="Welcome to STEM",
                        tip="Complete Chapter 1.",
                        guide_link="",
                    ),
                ],
            ),
        ],
        recommended_order=[
            "Start on Survival difficulty",
            "Follow missable guide for Chapter 3",
        ],
    )


def _make_result(output: LLMStrategyOutput | None = None) -> StrategyResult:
    return StrategyResult(
        app_id=268050,
        game_name="The Evil Within",
        model="claude-sonnet-4-6",
        output=output or _make_output(),
        created_at=datetime(2026, 6, 19, tzinfo=UTC),
    )


def _make_pending() -> list[PendingAchievement]:
    return [
        PendingAchievement(
            api_name="ACH_1",
            display_name="Don't Look Back",
            description="",
            is_hidden=False,
            schema_idx=1,
            icon_url="",
        ),
        PendingAchievement(
            api_name="ACH_2",
            display_name="Welcome to STEM",
            description="Complete Chapter 1",
            is_hidden=False,
            schema_idx=2,
            icon_url="",
        ),
    ]


def _docx_full_text(path: Path) -> str:
    doc = Document(str(path))
    return "\n".join(_para_full_text(p) for p in doc.paragraphs)


def _write_docx(tmp_path: Path, lines: list[str]) -> Path:
    doc = Document()
    for line in lines:
        doc.add_paragraph(line)
    path = tmp_path / "test.docx"
    doc.save(str(path))
    return path


# ── render_to_text ────────────────────────────────────────────────────────────


def test_render_to_text_includes_summary():
    output = _make_output()
    assert output.summary in render_to_text(output)


def test_render_to_text_includes_runs_and_hours():
    output = _make_output()
    text = render_to_text(output)
    assert str(output.total_runs) in text
    assert output.estimated_hours in text


def test_render_to_text_includes_section_titles():
    output = _make_output()
    text = render_to_text(output)
    for section in output.sections:
        assert section.title in text


def test_render_to_text_includes_all_display_names():
    output = _make_output()
    text = render_to_text(output)
    for section in output.sections:
        for item in section.items:
            assert item.display_name in text


def test_render_to_text_includes_all_tips():
    output = _make_output()
    text = render_to_text(output)
    for section in output.sections:
        for item in section.items:
            assert item.tip in text


def test_render_to_text_includes_recommended_order():
    output = _make_output()
    text = render_to_text(output)
    for step in output.recommended_order:
        assert step in text


def test_render_to_text_includes_my_notes_heading():
    assert "My Notes" in render_to_text(_make_output())


def test_render_to_text_checkbox_prefix_on_items():
    text = render_to_text(_make_output())
    assert "☐  Don't Look Back" in text


# ── render_docx ───────────────────────────────────────────────────────────────


def test_render_docx_creates_file(tmp_path):
    path = tmp_path / "guide.docx"
    render_docx(_make_result(), _make_pending(), path)
    assert path.exists()


def test_render_docx_contains_game_name(tmp_path):
    path = tmp_path / "guide.docx"
    render_docx(_make_result(), _make_pending(), path)
    assert "The Evil Within" in _docx_full_text(path)


def test_render_docx_contains_summary(tmp_path):
    path = tmp_path / "guide.docx"
    result = _make_result()
    render_docx(result, _make_pending(), path)
    assert result.output.summary in _docx_full_text(path)


def test_render_docx_contains_all_display_names(tmp_path):
    path = tmp_path / "guide.docx"
    result = _make_result()
    render_docx(result, _make_pending(), path)
    full_text = _docx_full_text(path)
    for section in result.output.sections:
        for item in section.items:
            assert item.display_name in full_text


def test_render_docx_contains_all_tips(tmp_path):
    path = tmp_path / "guide.docx"
    result = _make_result()
    render_docx(result, _make_pending(), path)
    full_text = _docx_full_text(path)
    for section in result.output.sections:
        for item in section.items:
            assert item.tip in full_text


def test_render_docx_contains_recommended_order(tmp_path):
    path = tmp_path / "guide.docx"
    result = _make_result()
    render_docx(result, _make_pending(), path)
    full_text = _docx_full_text(path)
    for step in result.output.recommended_order:
        assert step in full_text


def test_render_docx_pending_count_in_meta(tmp_path):
    path = tmp_path / "guide.docx"
    pending = _make_pending()
    render_docx(_make_result(), pending, path)
    assert f"Pending: {len(pending)}" in _docx_full_text(path)


def test_render_docx_my_notes_section_present(tmp_path):
    path = tmp_path / "guide.docx"
    render_docx(_make_result(), _make_pending(), path)
    assert "My Notes" in _docx_full_text(path)


# ── extract_annotations ───────────────────────────────────────────────────────


def test_extract_annotations_no_changes(tmp_path):
    baseline = "Line one\nLine two\nLine three"
    path = _write_docx(tmp_path, ["Line one", "Line two", "Line three"])
    result = extract_annotations(path, baseline)
    assert result.added_lines == []
    assert result.deleted_lines == []
    assert result.modified_lines == []


def test_extract_annotations_detects_added_line(tmp_path):
    baseline = "Line one\nLine two"
    path = _write_docx(tmp_path, ["Line one", "Line two", "New user note"])
    result = extract_annotations(path, baseline)
    assert "New user note" in result.added_lines


def test_extract_annotations_detects_deleted_line(tmp_path):
    baseline = "Line one\nLine two\nLine three"
    path = _write_docx(tmp_path, ["Line one", "Line three"])
    result = extract_annotations(path, baseline)
    assert "Line two" in result.deleted_lines


def test_extract_annotations_detects_modified_line(tmp_path):
    baseline = "Line one\nOriginal tip here"
    path = _write_docx(tmp_path, ["Line one", "Modified tip here"])
    result = extract_annotations(path, baseline)
    assert "Modified tip here" in result.modified_lines


def test_extract_annotations_returns_docx_path(tmp_path):
    path = _write_docx(tmp_path, ["some text"])
    result = extract_annotations(path, "some text")
    assert result.docx_path == str(path)


def test_extract_annotations_ignores_blank_paragraphs(tmp_path):
    baseline = "Line one\nLine two"
    doc = Document()
    doc.add_paragraph("Line one")
    doc.add_paragraph("")
    doc.add_paragraph("Line two")
    path = tmp_path / "blank.docx"
    doc.save(str(path))
    result = extract_annotations(path, baseline)
    assert result.added_lines == []
    assert result.deleted_lines == []
